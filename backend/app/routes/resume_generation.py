from flask_jwt_extended import jwt_required, get_jwt_identity
from flask import Blueprint, request, jsonify, send_file
import requests
import os
import json
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from bs4 import BeautifulSoup
import boto3
from datetime import datetime

resume_bp = Blueprint("resume_generation", __name__)

OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY")
OPENAI_URL = "https://api.openai.com/v1/chat/completions"

# AWS S3 Configuration
AWS_ACCESS_KEY = os.environ.get("AWS_ACCESS_KEY_ID")
AWS_SECRET_KEY = os.environ.get("AWS_SECRET_ACCESS_KEY")
AWS_S3_BUCKET = os.environ.get("S3_BUCKET_NAME")
AWS_REGION = os.environ.get("AWS_REGION", "us-east-1")

s3_client = boto3.client(
    's3',
    aws_access_key_id=AWS_ACCESS_KEY,
    aws_secret_access_key=AWS_SECRET_KEY,
    region_name=AWS_REGION
)


@resume_bp.route("/api/generate_resume", methods=['POST'])
def generate_resume():
    if not OPENAI_API_KEY:
        return jsonify({"error": "OPENAI_API_KEY not set"}), 500

    form_data = request.get_json()
    if not form_data:
        return jsonify({"error": "No form data provided"}), 400

    # Professional Resume Design - Exact Template Match
    system_prompt = """
You are a resume HTML generator. Generate ONLY pure HTML code with embedded CSS.

CRITICAL REQUIREMENTS:
- Output ONLY HTML, nothing else
- Start with <!DOCTYPE html>
- Include <html>, <head>, <body> tags
- Put all CSS in <style> tag in <head>
- NO markdown, NO backticks, NO code fences
- NO explanations before or after
- White background (#ffffff)
- Black text (#000000)
- Arial font
- Professional traditional resume layout
- Left AND RIGHT margin/padding: 40-50px for indentation
- Font size: 11-12px for body text, slightly larger for headers
- ONLY INCLUDE SECTIONS THAT HAVE DATA - skip empty sections

EXACT LAYOUT PATTERN:
1. NAME (large 20px, bold, centered)
2. Location • Email (centered, 11px)
3. Phone • LinkedIn URL (centered, 11px)
4. BLANK LINE

SECTION HEADERS (ONLY capitalize first letter, NOT all caps):
Education
Relevant Experience
Additional Experience
Community Involvement
Projects
Skills Summary
Certifications

STRUCTURE FOR EACH SECTION:
- Section Title: Title Case (not ALL CAPS), bold, followed by horizontal line (full width, 1px solid black)
- Company/School/Project: **Bold**, City, State (or project details) - format with HTML <b> tags, NO asterisks
- Position/Degree/Project Title: Position title, Duration/Dates
- Bullets: · Detailed bullet point (2-3 lines describing responsibilities and achievements)
- Spacing: Single line between entries, blank line before next section
- Left padding: 40px
- Right padding: 40px
- Total body width with padding: 816px (8.5 inches)

PROJECTS SECTION FORMAT (if projects exist):
Project Name (bold, no asterisks)
Project description or link
· Achievement/detail about project
· Technology used / impact

SKILLS SECTION FORMAT:
Technical Skills: item1, item2, item3
Social Media: item1, item2, item3
Soft Skills: item1, item2, item3
Languages: item1, item2

CSS REQUIREMENTS:
- body: background white, color black, font Arial 11pt, max-width 816px
- Page width: 8.5 x 11 inches (816px)
- Margins: 0.5 inch all sides (40px left and right)
- Left AND right padding on all content: 40px
- Section headers: bold, Title Case, with horizontal line below (1px solid black)
- Company/School/Project names: bold
- No colors, pure black and white
- Line spacing: 1.4 for body content
- Horizontal lines: full width, 1px solid black
- Hide/skip sections with no data

OUTPUT ONLY THE HTML CODE STARTING WITH <!DOCTYPE html>
"""

    user_prompt = f"""
Generate a professional resume in HTML/CSS matching this EXACT layout:

NAME (centered, large bold)
Location, State • email@domain.com (centered)
Phone • LinkedIn URL (centered)

[ONLY INCLUDE SECTIONS BELOW IF THEY HAVE DATA - SKIP EMPTY SECTIONS]

Education
[with horizontal line below]

Relevant Experience
[with horizontal line below]

Additional Experience
[with horizontal line below]

Community Involvement
[with horizontal line below]

Projects
[with horizontal line below]

Skills Summary
[Technical Skills, Social Media, Soft Skills, Languages]

Certifications
[with horizontal line below]

Data to use:
{json.dumps(form_data, indent=2)}

IMPORTANT RULES:
1. Only display sections that have actual data
2. Skip/hide sections with no data (e.g., if no certifications, don't show certifications section)
3. Include Projects section if projects exist
4. Create detailed, professional bullet points for EVERY job and project

JOB DESCRIPTION GENERATION:
Create detailed, professional bullet points for EVERY job entry. If description is minimal or empty:

For specific job titles, generate relevant bullets:
- "Data Analyst" → "• Analyzed large datasets using SQL and Python to identify business trends and opportunities for improvement", "• Created comprehensive reports and visualizations to present findings to senior management", "• Collaborated with cross-functional teams to implement data-driven solutions improving efficiency by 15%"

- "Software Developer" → "• Designed and developed full-stack web applications using modern technologies and best practices", "• Debugged complex issues and optimized code performance, reducing load times by 20%", "• Participated in code reviews and mentored junior developers on coding standards and design patterns"

- "Marketing Intern" → "• Assisted in planning and executing integrated marketing campaigns across multiple channels", "• Conducted market research and competitor analysis to inform marketing strategy", "• Created engaging content for social media and marketing materials, increasing engagement by 25%"

- "Sales Associate" → "• Provided excellent customer service and product expertise to drive sales and customer satisfaction", "• Exceeded monthly sales targets consistently through effective upselling and relationship building", "• Processed transactions accurately and maintained store displays to ensure positive shopping experience"

PROJECTS GENERATION:
For each project, include:
- **Project Name** (bold)
- Brief description or technologies used
- 2-3 bullet points with achievements and impact

Requirements:
1. White background, black text, Arial font, 11-12pt
2. Centered name at top (large, bold ~20px)
3. Contact info centered below name (location • email, then phone • linkedin)
4. Section headers: Title Case (not ALL CAPS) with horizontal line below
5. Format: Company/School/Project Name in bold (use <b> HTML tags), City, State - NO ASTERISKS
6. Position/degree/project details on second line
7. 3-4 detailed bullet points with · symbol for each entry
8. Skills: "Technical Skills: x, y, z" format
9. Left AND right padding on all content: 40px
10. One page layout
11. ONLY INCLUDE SECTIONS WITH DATA - skip empty sections
12. Include Projects section if projects exist
13. Project names should be bold with NO asterisks (use HTML <b> tags)
14. Company names AND locations should be bold

Output ONLY the complete HTML code starting with <!DOCTYPE html>
Do NOT include backticks, markdown, explanations, or code fences.
"""

    headers = {
        "Authorization": f"Bearer {OPENAI_API_KEY}",
        "Content-Type": "application/json"
    }

    payload = {
        "model": "gpt-4o-mini",
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ],
        "temperature": 0.6,
        "max_tokens": 2000,
        "stream": False
    }

    try:
        res = requests.post(OPENAI_URL, headers=headers, json=payload)
        res.raise_for_status()
        data = res.json()

        ai_text = data.get("choices", [{}])[0].get("message", {}).get("content", "")

        if not ai_text:
            return jsonify({
                "error": "No resume text returned from AI",
                "details": data
            }), 500

        return jsonify({
            "message": "Resume generated successfully",
            "resume_text": ai_text
        }), 200

    except Exception as e:
        return jsonify({
            "error": "Failed to generate resume",
            "details": str(e)
        }), 500


def html_to_docx(html_content):
    """Convert resume HTML to DOCX document"""
    doc = Document()
    doc.core_properties.author = "HireHub Resume Builder"
    
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Remove script and style tags
    for script in soup(["script", "style"]):
        script.decompose()
    
    # Get all text and structure from the entire document
    def process_element(element, doc):
        """Recursively process elements and add to document"""
        if isinstance(element, str):
            text = element.strip()
            if text and len(text) > 1:
                doc.add_paragraph(text)
            return
        
        if not hasattr(element, 'name'):
            return
        
        text = element.get_text(strip=True) if hasattr(element, 'get_text') else ''
        
        if not text or len(text) < 2:
            # Still process children even if this element has no text
            if element.name in ['div', 'span', 'section', 'article']:
                for child in element.children:
                    process_element(child, doc)
            return
        
        # Handle headers
        if element.name == 'h1':
            p = doc.add_paragraph(text, style='Heading 1')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif element.name == 'h2':
            doc.add_paragraph(text, style='Heading 2')
        elif element.name == 'h3':
            doc.add_paragraph(text, style='Heading 3')
        elif element.name == 'h4':
            doc.add_paragraph(text, style='Heading 4')
        
        # Handle lists
        elif element.name == 'ul' or element.name == 'ol':
            for li in element.find_all('li', recursive=False):
                li_text = li.get_text(strip=True)
                if li_text:
                    doc.add_paragraph(li_text, style='List Bullet')
        
        # Handle list items
        elif element.name == 'li':
            doc.add_paragraph(text, style='List Bullet')
        
        # Handle paragraphs, divs, and other text containers
        else:
            doc.add_paragraph(text)
    
    # Process the entire document
    for element in soup.find_all():
        if element.name and element.name not in ['html', 'head', 'body', 'meta', 'link', 'title']:
            # Only process elements that are direct children to avoid duplicates
            if not element.find_parent(['h1', 'h2', 'h3', 'h4', 'p', 'div', 'li', 'ul', 'ol']):
                process_element(element, doc)
    
    # If document is still empty, try to get all text
    if len(doc.paragraphs) == 0:
        all_text = soup.get_text(strip=True)
        if all_text:
            # Split by newlines and add as paragraphs
            for line in all_text.split('\n'):
                line = line.strip()
                if line and len(line) > 1:
                    doc.add_paragraph(line)
    
    return doc


@resume_bp.route("/api/generate-docx", methods=['POST'])
def generate_docx():
    """Convert resume HTML to DOCX format and download"""
    try:
        data = request.get_json()
        html_content = data.get('html', '')
        
        if not html_content:
            return jsonify({"error": "No HTML content provided"}), 400
        
        doc = html_to_docx(html_content)
        
        docx_io = BytesIO()
        doc.save(docx_io)
        docx_io.seek(0)
        
        return send_file(
            docx_io,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='resume.docx'
        )
        
    except Exception as e:
        print(f"Error generating DOCX: {str(e)}")
        return jsonify({
            "error": "Failed to generate DOCX",
            "details": str(e)
        }), 500


@resume_bp.route("/api/save-resume-to-storage", methods=['POST'])
@jwt_required()
def save_resume_to_storage():
    """Convert resume HTML to DOCX and save to AWS S3"""
    try:
        print("=== Starting save-resume-to-storage ===")
        print(f"AWS_ACCESS_KEY set: {bool(AWS_ACCESS_KEY)}")
        print(f"AWS_SECRET_KEY set: {bool(AWS_SECRET_KEY)}")
        print(f"AWS_S3_BUCKET: {AWS_S3_BUCKET}")
        
        if not all([AWS_ACCESS_KEY, AWS_SECRET_KEY, AWS_S3_BUCKET]):
            error_msg = "AWS configuration not set"
            print(f"Error: {error_msg}")
            return jsonify({"error": error_msg}), 500
        
        data = request.get_json()
        print(f"Request data keys: {data.keys() if data else 'None'}")
        
        html_content = data.get('html', '')
        user_id = data.get('user_id')
        filename = data.get('filename', 'resume.docx')
        
        print(f"HTML content length: {len(html_content)}")
        print(f"User ID: {user_id}")
        print(f"Filename: {filename}")
        
        if not html_content:
            error_msg = "No HTML content provided"
            print(f"Error: {error_msg}")
            return jsonify({"error": error_msg}), 400
        
        if not user_id:
            error_msg = "User ID is required"
            print(f"Error: {error_msg}")
            return jsonify({"error": error_msg}), 400
        
        print("Converting HTML to DOCX...")
        doc = html_to_docx(html_content)
        
        docx_io = BytesIO()
        doc.save(docx_io)
        docx_io.seek(0)
        print(f"DOCX file size: {len(docx_io.getvalue())} bytes")
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        s3_key = f"resumes/{user_id}/{timestamp}_{filename}"
        print(f"S3 Key: {s3_key}")
        
        print("Uploading to S3...")
        s3_client.put_object(
            Bucket=AWS_S3_BUCKET,
            Key=s3_key,
            Body=docx_io.getvalue(),
            ContentType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        print("Successfully uploaded to S3")
        
        from app.extensions import db
        from app.models.document import Document
        from app.models.user import User
        
        try:
            current_user_id = get_jwt_identity()
            user = User.query.get(current_user_id)
            
            if not user:
                print(f"Warning: User not found for ID: {current_user_id}")
                return jsonify({
                    "message": "Resume saved to S3 but could not link to user account",
                    "s3_key": s3_key,
                    "filename": filename
                }), 200
            
            new_document = Document(
                user_email=user.email,
                file_path=s3_key,
                original_filename=filename,
                document_type='resume'
            )
            db.session.add(new_document)
            db.session.commit()
            print(f"Database record created for document ID: {new_document.document_id}")
            
        except Exception as db_error:
            db.session.rollback()
            print(f"Warning: Could not create database record: {str(db_error)}")
            import traceback
            traceback.print_exc()
        
        return jsonify({
            "message": "Resume saved to storage successfully",
            "s3_key": s3_key,
            "filename": filename
        }), 200
        
    except Exception as e:
        error_msg = f"Error saving resume to storage: {str(e)}"
        print(error_msg)
        import traceback
        traceback.print_exc()
        return jsonify({
            "error": "Failed to save resume to storage",
            "details": str(e)
        }), 500